#!/usr/bin/env python3
"""Convert primaria-secundaria-final.html to a Word (.docx) document.

Usage:
    python3 scripts/html_to_docx.py
"""

import re
from pathlib import Path
from bs4 import BeautifulSoup, NavigableString, Tag
from docx import Document
from docx.shared import Pt, Inches, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.style import WD_STYLE_TYPE
from docx.oxml.ns import qn, nsdecls
from docx.oxml import parse_xml

ROOT = Path(__file__).resolve().parent.parent
HTML_PATH = ROOT / "documento" / "primaria-secundaria-final.html"
DOCX_PATH = ROOT / "documento" / "primaria-secundaria-final.docx"


def setup_styles(doc):
    """Create custom styles for the document."""
    styles = doc.styles

    # Blockquote style
    bq = styles.add_style("Blockquote", WD_STYLE_TYPE.PARAGRAPH)
    bq.font.italic = True
    bq.font.size = Pt(10.5)
    bq.font.color.rgb = RGBColor(0x33, 0x33, 0x33)
    bq.paragraph_format.left_indent = Cm(1.5)
    bq.paragraph_format.space_before = Pt(6)
    bq.paragraph_format.space_after = Pt(2)

    # Citation style (for blockquote attribution)
    cite = styles.add_style("Citation", WD_STYLE_TYPE.PARAGRAPH)
    cite.font.size = Pt(9.5)
    cite.font.color.rgb = RGBColor(0x55, 0x55, 0x55)
    cite.paragraph_format.left_indent = Cm(1.5)
    cite.paragraph_format.space_before = Pt(0)
    cite.paragraph_format.space_after = Pt(8)

    # Cover subtitle
    sub = styles.add_style("CoverSubtitle", WD_STYLE_TYPE.PARAGRAPH)
    sub.font.size = Pt(14)
    sub.font.color.rgb = RGBColor(0x66, 0x66, 0x66)
    sub.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
    sub.paragraph_format.space_after = Pt(4)

    # Cover meta
    meta = styles.add_style("CoverMeta", WD_STYLE_TYPE.PARAGRAPH)
    meta.font.size = Pt(10)
    meta.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
    meta.paragraph_format.space_before = Pt(2)
    meta.paragraph_format.space_after = Pt(2)

    return styles


def decode_html(text):
    """Clean up HTML entities and whitespace."""
    if not text:
        return ""
    import html as htmlmod
    text = htmlmod.unescape(text)
    text = re.sub(r"\s+", " ", text)
    return text.strip()


def get_text(el):
    """Get all text from an element, recursively."""
    return decode_html(el.get_text())


def add_inline(run, el):
    """Apply inline formatting to a run based on its parent tags."""
    if not isinstance(el, Tag):
        return
    parents = [p.name for p in el.parents if isinstance(p, Tag)]
    if el.name in ("strong", "b") or "strong" in parents or "b" in parents:
        run.bold = True
    if el.name in ("em", "i") or "em" in parents or "i" in parents:
        run.italic = True


def add_rich_paragraph(doc, element, style=None):
    """Add a paragraph with mixed bold/italic/link formatting."""
    para = doc.add_paragraph(style=style)
    for child in element.children:
        if isinstance(child, NavigableString):
            text = decode_html(str(child))
            if text:
                para.add_run(text)
        elif isinstance(child, Tag):
            if child.name == "br":
                para.add_run("\n")
            elif child.name in ("strong", "b"):
                run = para.add_run(get_text(child))
                run.bold = True
            elif child.name in ("em", "i"):
                run = para.add_run(get_text(child))
                run.italic = True
            elif child.name == "a":
                run = para.add_run(get_text(child))
                run.underline = True
                run.font.color.rgb = RGBColor(0x15, 0x65, 0xC0)
            else:
                # Recurse into nested elements
                for sub in child.children:
                    if isinstance(sub, NavigableString):
                        text = decode_html(str(sub))
                        if text:
                            run = para.add_run(text)
                            add_inline(run, child)
                    elif isinstance(sub, Tag):
                        run = para.add_run(get_text(sub))
                        if sub.name in ("strong", "b"):
                            run.bold = True
                        elif sub.name in ("em", "i"):
                            run.italic = True
                        add_inline(run, child)
    return para


def add_table(doc, table_el):
    """Convert an HTML table to a Word table."""
    rows_data = []

    thead = table_el.find("thead")
    tbody = table_el.find("tbody")

    header_rows = []
    if thead:
        for tr in thead.find_all("tr"):
            cells = [get_text(c) for c in tr.find_all(["th", "td"])]
            header_rows.append(cells)

    body_rows = []
    source = tbody if tbody else table_el
    for tr in source.find_all("tr", recursive=False if tbody else True):
        if thead and tr.parent == thead:
            continue
        cells = [get_text(c) for c in tr.find_all(["th", "td"])]
        body_rows.append(cells)

    rows_data = header_rows + body_rows
    if not rows_data:
        return

    num_cols = max(len(r) for r in rows_data)
    # Pad short rows
    for r in rows_data:
        while len(r) < num_cols:
            r.append("")

    table = doc.add_table(rows=len(rows_data), cols=num_cols)
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER

    for i, row_data in enumerate(rows_data):
        row = table.rows[i]
        for j, cell_text in enumerate(row_data):
            cell = row.cells[j]
            cell.text = cell_text
            for paragraph in cell.paragraphs:
                paragraph.paragraph_format.space_before = Pt(2)
                paragraph.paragraph_format.space_after = Pt(2)
                for run in paragraph.runs:
                    run.font.size = Pt(9.5)

    # Bold header rows
    for i in range(len(header_rows)):
        for cell in table.rows[i].cells:
            for paragraph in cell.paragraphs:
                for run in paragraph.runs:
                    run.bold = True
            # Shade header
            shading = parse_xml(
                f'<w:shd {nsdecls("w")} w:fill="E8EAF6" w:val="clear"/>'
            )
            cell._tc.get_or_add_tcPr().append(shading)

    doc.add_paragraph("")  # spacing after table


def add_blockquote(doc, bq_el):
    """Convert an HTML blockquote to styled paragraphs."""
    for child in bq_el.children:
        if isinstance(child, Tag):
            if child.name == "p":
                add_rich_paragraph(doc, child, style="Blockquote")
            elif child.name == "cite":
                p = doc.add_paragraph(get_text(child), style="Citation")


def add_list(doc, ul_el):
    """Convert an HTML list to Word list items."""
    for li in ul_el.find_all("li", recursive=False):
        add_rich_paragraph(doc, li, style="List Bullet")


def process_cover(doc, section):
    """Process the cover page section."""
    subtitle = section.find("p", class_="cover-subtitle")
    if subtitle:
        doc.add_paragraph(get_text(subtitle), style="CoverSubtitle")

    title = section.find("h1")
    if title:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(get_text(title))
        run.font.size = Pt(28)
        run.bold = True
        run.font.color.rgb = RGBColor(0x1A, 0x23, 0x7E)

    level = section.find("p", class_="cover-level")
    if level:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(get_text(level))
        run.font.size = Pt(16)
        run.font.color.rgb = RGBColor(0x33, 0x33, 0x33)

    doc.add_paragraph("")  # spacer

    meta_div = section.find("div", class_="cover-meta")
    if meta_div:
        for p_tag in meta_div.find_all("p"):
            add_rich_paragraph(doc, p_tag, style="CoverMeta")

    doc.add_page_break()


def process_toc(doc, nav):
    """Process table of contents as a simple list."""
    doc.add_heading("Tabla de Contenidos", level=1)
    ol = nav.find("ol")
    if ol:
        for i, li in enumerate(ol.find_all("li", recursive=False), 1):
            p = doc.add_paragraph()
            run = p.add_run(f"{i}. {get_text(li)}")
            run.font.size = Pt(11)
    doc.add_page_break()


def process_section(doc, section):
    """Process a regular content section."""
    for child in section.children:
        if isinstance(child, NavigableString):
            continue
        if not isinstance(child, Tag):
            continue

        if child.name == "h2":
            doc.add_heading(get_text(child), level=1)
        elif child.name == "h3":
            doc.add_heading(get_text(child), level=2)
        elif child.name == "h4":
            doc.add_heading(get_text(child), level=3)
        elif child.name == "p":
            cls = child.get("class", [])
            if "demo-link" in str(cls):
                continue  # skip "Explora en 3D" links
            add_rich_paragraph(doc, child)
        elif child.name == "a" and "demo-link" in (child.get("class") or []):
            continue
        elif child.name == "table":
            add_table(doc, child)
        elif child.name == "blockquote":
            add_blockquote(doc, child)
        elif child.name in ("ul", "ol"):
            add_list(doc, child)
        elif child.name == "div":
            # Handle perfil-cards and entrevista blocks
            cls = " ".join(child.get("class", []))
            if "perfil-card" in cls or "entrevista" in cls:
                # Add heading if present
                h3 = child.find("h3")
                if h3:
                    doc.add_heading(get_text(h3), level=2)
                # Process inner content
                inner = child.find("div", class_="entrevista-respuesta") or child
                for el in inner.children:
                    if isinstance(el, Tag):
                        if el.name == "h3":
                            continue  # already handled
                        elif el.name == "p":
                            add_rich_paragraph(doc, el)
                        elif el.name == "blockquote":
                            add_blockquote(doc, el)
                        elif el.name == "table":
                            add_table(doc, el)
                        elif el.name in ("ul", "ol"):
                            add_list(doc, el)
                        elif el.name == "div":
                            # Nested div (like Sofia's diary entry)
                            for sub_el in el.children:
                                if isinstance(sub_el, Tag):
                                    if sub_el.name == "p":
                                        add_rich_paragraph(doc, sub_el)
            elif "entrevista-respuesta" in cls:
                for el in child.children:
                    if isinstance(el, Tag):
                        if el.name == "p":
                            add_rich_paragraph(doc, el)


def main():
    html = HTML_PATH.read_text(encoding="utf-8")
    soup = BeautifulSoup(html, "html.parser")

    doc = Document()

    # Page setup
    for section in doc.sections:
        section.top_margin = Cm(2.5)
        section.bottom_margin = Cm(2.5)
        section.left_margin = Cm(2.5)
        section.right_margin = Cm(2.5)

    # Default font
    style = doc.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(11)
    style.paragraph_format.space_after = Pt(6)

    # Heading styles
    for level in range(1, 4):
        h_style = doc.styles[f"Heading {level}"]
        h_style.font.name = "Calibri"
        h_style.font.color.rgb = RGBColor(0x1A, 0x23, 0x7E)

    setup_styles(doc)

    article = soup.find("article", class_="document")
    if not article:
        print("ERROR: No <article class='document'> found.")
        return

    for section_el in article.children:
        if not isinstance(section_el, Tag):
            continue

        section_id = section_el.get("id", "")
        section_cls = " ".join(section_el.get("class", []))

        if "cover" in section_cls:
            process_cover(doc, section_el)
        elif section_el.name == "nav" and "toc" in section_cls:
            process_toc(doc, section_el)
        elif section_el.name == "section":
            process_section(doc, section_el)

    doc.save(str(DOCX_PATH))
    print(f"OK: {DOCX_PATH}")
    print(f"    {DOCX_PATH.stat().st_size:,} bytes")


if __name__ == "__main__":
    main()
